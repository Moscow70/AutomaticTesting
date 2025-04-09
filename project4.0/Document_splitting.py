import re
from docx import Document
from langchain_community.document_loaders import PyPDFLoader
import pdfplumber

class DocumentProcessor:
    def __init__(self, max_length=2048, overlap_length=512):
        self.max_length = max_length
        self.overlap_length = overlap_length

    def read_txt_file(self, file_path):
        """读取纯文本文件的内容"""
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()
        return content

    def read_docx_file(self, file_path):
        """读取Word文档的内容，包括文本和表格"""
        doc = Document(file_path)
        content = []
        for para in doc.paragraphs:
            content.append(para.text)
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text)
                table_content.append(" | ".join(row_content))
            content.append("\n".join(table_content))
        return "\n\n".join(content)

    def read_pdf_file(self, file_path):
        """读取PDF文件的内容，包括文本和表格"""
        # 读取PDF文本内容
        loader = PyPDFLoader(file_path)
        pages = loader.load_and_split()
        text_content = []
        for page in pages:
            text = page.page_content
            text_content.append(text)
        
        # 读取PDF表格内容
        table_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    table_str = "\n".join([" | ".join(map(str, row)) for row in table])
                    table_content.append(table_str)
        
        return "\n\n".join(text_content + table_content)

    def detect_tables_in_text(self, text):
        """检测文本中的表格，并返回表格的起始和结束位置"""
        table_pattern = re.compile(r'(\|.*\|\n\|.*\|.*\|[\s\S]*?\n\|.*\|)')
        tables = []
        for match in table_pattern.finditer(text):
            tables.append((match.start(), match.end()))
        return tables

    def split_document_with_tables(self, text):
        """将文档按最大长度切分，并保留重叠部分，同时确保表格的完整性"""
        parts = []
        start = 0
        while start < len(text):
            end = min(start + self.max_length, len(text))
            
            # 检查是否有表格跨越当前切片
            for table_start, table_end in self.detect_tables_in_text(text):
                if start <= table_start < end and table_end > end:
                    end = table_end
            
            part = text[start:end]
            parts.append(part)

            if end == len(text):
                start = end
            else:
                start = end - self.overlap_length
        
        return parts

    def process_document(self, file_path):
        """主函数，根据文件类型读取内容并切割"""
        if file_path.endswith('.txt'):
            content = self.read_txt_file(file_path)
        elif file_path.endswith('.docx'):
            content = self.read_docx_file(file_path)
        elif file_path.endswith('.pdf'):
            content = self.read_pdf_file(file_path)
        else:
            raise ValueError("Unsupported file type")

        # 切割文档
        parts = self.split_document_with_tables(content)
        return parts