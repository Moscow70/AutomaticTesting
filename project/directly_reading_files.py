from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document

load_dotenv(dotenv_path="../.env", override=True)

api_key = os.environ.get("DEEPSEEK_API_KEY")
base_url = os.environ.get("DEEPSEEK_BASE_URL")

class reading_files_and_answering:

    def __init__(self, file_path):
        self.file_path = file_path
        self.client = OpenAI(
            api_key = api_key,
            base_url = base_url
        )

    # llm = OpenAI(
    #     api_key = api_key,
    #     base_url = base_url
    # )

    def get_case(self, prompt, model):
        # 读取文件内容
        # with open(self.file_path, 'r', encoding='utf-8') as file:
        #     content = file.read()
        doc = Document(self.file_path)

        content = []

        for para in doc.paragraphs:
            content.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content.append(cell.text)
        
        content_str = '\n'.join(content)
        # prompt = content + '\n' + prompt + '\n'

        chat_completions = self.client.chat.completions.create(
            messages = [
                # {"role": "system", "content": "你是一个软件测试工程师，你需要根据读入的文档生成对应的接口测试代码，要全面地覆盖文档中出现的情况"},
                {"role": "system", "content": prompt},
                {"role": "user", "content": content_str}
            ],
            model= model,
        )
        output = chat_completions.choices[0].message.content
        return output

    # question = "根据提示词，按照格式生成对应的接口测试代码，使用python语言的格式，文字解释的部分每一行前面都要加上#。"
    # file_path = "D:/workspace/originalfiles/使用手册.txt"

    # with open(file_path, 'r', encoding='utf-8') as file:
    #     content = file.read()

    # prompt = content + '\n' + question + '\n'

    # answer = get_case(prompt)

    # answer_path = "D:/workspace/output/answer1.txt"

    # with open(answer_path, 'w', encoding='utf-8') as file:
    #     file.write(answer)




